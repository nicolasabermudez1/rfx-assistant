"""Spec Builder tab.

Chatbot asks the user what they want to buy. Gemini (or product-specific
fallbacks) generates a structured spec table for any spend category. Team
members defined at the welcome screen own the rows; either can send email
reminders to the other.
"""
from __future__ import annotations

import time
from datetime import datetime, timezone

import pandas as pd
import streamlit as st

from rfx_assistant import agents
from rfx_assistant.ui import team

# ---------------------------------------------------------------------------
# State helpers
# ---------------------------------------------------------------------------

def _init():
    st.session_state.setdefault("sb_messages", [])
    st.session_state.setdefault("sb_spec", None)
    st.session_state.setdefault("sb_stage", 0)
    st.session_state.setdefault("sb_activity", [])
    st.session_state.setdefault("sb_reminder_drafted", {})
    st.session_state.setdefault("sb_reminder_sent", {})
    st.session_state.setdefault("sb_generating", False)
    # Dynamic, category-tailored clarifying questions (populated after stage 0)
    st.session_state.setdefault("sb_dynamic_questions", [])
    st.session_state.setdefault("sb_category_text", "")


def _active_user() -> dict | None:
    # Always the logged-in user — no persona switching.
    return team.get_me()


def _push_activity(action: str, actor: dict, icon: str = "✏️"):
    st.session_state.sb_activity.insert(0, {
        "t": datetime.now(tz=timezone.utc).isoformat(),
        "actor": actor.get("name", "?") if actor else "?",
        "icon": icon,
        "action": action,
    })


def _human_time(iso: str) -> str:
    try:
        t = datetime.fromisoformat(iso)
        secs = (datetime.now(tz=timezone.utc) - t).total_seconds()
        if secs < 60:
            return "just now"
        if secs < 3600:
            return f"{int(secs // 60)}m ago"
        return f"{int(secs // 3600)}h ago"
    except Exception:
        return iso


# ---------------------------------------------------------------------------
# Chatbot
# ---------------------------------------------------------------------------

_GREETING = (
    "Hi! I'm your RFx Assistant. **What are you looking to procure?** "
    "Describe it in a few words — it could be anything from software licences "
    "to industrial batteries to professional services.\n\n"
    "💡 *Tip: drop existing specs, RFP templates, or any reference document "
    "into the upload zone below and I'll fold them into the spec.*"
)


# ---------------------------------------------------------------------------
# File upload — extract plain text from any document the user attaches
# ---------------------------------------------------------------------------

_SUPPORTED_TYPES = ["pdf", "docx", "xlsx", "xlsm", "csv", "txt", "md", "json", "yml", "yaml"]


def _extract_text(uploaded_file) -> str:
    """Best-effort plain-text extraction from an uploaded file."""
    name = (uploaded_file.name or "").lower()
    try:
        uploaded_file.seek(0)
        if name.endswith(".pdf"):
            import pypdf
            reader = pypdf.PdfReader(uploaded_file)
            return "\n".join((p.extract_text() or "") for p in reader.pages)
        if name.endswith(".docx"):
            from docx import Document
            doc = Document(uploaded_file)
            parts = [p.text for p in doc.paragraphs if p.text]
            for tbl in doc.tables:
                for row in tbl.rows:
                    parts.append(" | ".join(cell.text for cell in row.cells))
            return "\n".join(parts)
        if name.endswith((".xlsx", ".xlsm", ".xls")):
            import pandas as pd
            sheets = pd.read_excel(uploaded_file, sheet_name=None)
            return "\n\n".join(
                f"## Sheet: {sn}\n{df.to_csv(index=False)}"
                for sn, df in sheets.items()
            )
        if name.endswith(".csv"):
            import pandas as pd
            df = pd.read_csv(uploaded_file)
            return df.to_csv(index=False)
        # Plain text formats (.txt, .md, .json, .yml, .yaml, unknown)
        data = uploaded_file.read()
        if isinstance(data, bytes):
            return data.decode("utf-8", errors="ignore")
        return str(data)
    except Exception as e:
        return f"[Could not extract text from {uploaded_file.name}: {e}]"


def _collect_attached_text(uploaded_files) -> str:
    """Extract + cache text from every attached file. Returns concatenated text."""
    if not uploaded_files:
        return ""
    cache = st.session_state.setdefault("sb_extract_cache", {})
    parts = []
    for f in uploaded_files:
        key = f"{f.name}|{getattr(f, 'size', '?')}"
        if key not in cache:
            cache[key] = _extract_text(f)
        parts.append(f"\n\n--- FILE: {f.name} ---\n{cache[key]}")
    return "\n".join(parts)


def _bot_reply_for_stage(stage: int, user_msg: str) -> tuple[str, bool]:
    """Bot reply for the current stage. Uses category-aware dynamic questions
    generated after the user answers stage 0."""
    qs = st.session_state.get("sb_dynamic_questions", [])

    if stage == 0:
        # The user just told us what they want to buy. We've already kicked off
        # question generation in _render_chat — qs should be populated.
        category = st.session_state.get("sb_category_text", user_msg.strip())
        if qs:
            return (
                f"Got it — **{category}**. Two quick questions tailored to that:\n\n"
                f"**{qs[0]}**",
                False,
            )
        return (
            f"Got it — **{category}**. What's the approximate scale or scope?",
            False,
        )

    if stage == 1:
        if len(qs) >= 2:
            return (
                f"Thanks. Last one:\n\n**{qs[1]}** — or type **none** to proceed.",
                False,
            )
        return (
            "Thanks. Any key constraints, preferences, or compliance "
            "requirements? — or type **none** to proceed.",
            False,
        )

    return (
        "Perfect — I have everything I need. "
        "**Building your product-specific specification and scoring matrix now…**",
        True,
    )


# ---------------------------------------------------------------------------
# Public render
# ---------------------------------------------------------------------------

def render():
    _init()
    me = _active_user()
    if not me:
        st.warning("Please set up your profile on the welcome screen first.")
        return

    st.markdown("### Spec Builder")
    st.caption(
        "Describe what you're buying — the AI builds a product-specific spec "
        "table. Edit rows, assign owners, and send reminders to teammates."
    )

    st.divider()

    if st.session_state.sb_stage < 4:
        _render_chat()
    else:
        _render_workspace()


# ---------------------------------------------------------------------------
# Chat phase
# ---------------------------------------------------------------------------

def _render_chat():
    msgs = st.session_state.sb_messages

    if not msgs:
        msgs.append({"role": "assistant", "content": _GREETING})

    for m in msgs:
        with st.chat_message(m["role"]):
            st.markdown(m["content"])

    # ------------------------------------------------------------------
    # File uploader — accepts any reference docs the user wants the AI to use
    # ------------------------------------------------------------------
    uploaded = st.file_uploader(
        "📎  Attach existing specs, RFP templates, or any reference documentation "
        "(PDF, Word, Excel, CSV, TXT, MD, JSON, YAML — multiple files OK)",
        type=_SUPPORTED_TYPES,
        accept_multiple_files=True,
        key="sb_upload_widget",
    )
    if uploaded:
        names_html = " ".join(
            f"<span style='background:var(--surface-2);"
            f"border:1px solid var(--border);border-radius:12px;"
            f"padding:3px 10px;font-size:12px;margin-right:4px;"
            f"display:inline-block'>📄 {f.name}</span>"
            for f in uploaded
        )
        st.markdown(
            f"<div style='margin:6px 0 10px 0'>"
            f"<b>{len(uploaded)} file(s) attached:</b> {names_html}<br>"
            f"<span style='color:var(--muted);font-size:12px'>"
            f"The AI will use these to tailor the spec when you finish the questions."
            f"</span></div>",
            unsafe_allow_html=True,
        )

    if st.session_state.sb_stage == 3 and not st.session_state.sb_generating:
        st.session_state.sb_generating = True
        attached_text = _collect_attached_text(uploaded)
        live = agents.gemini_key_available()
        msg = "Calling Gemini to tailor your spec + scoring matrix…" if live else \
              "Building your spec from product-specific templates…"
        if attached_text:
            msg = msg.rstrip("…") + " (using your uploaded docs)…"
        with st.spinner(msg):
            spec = agents.generate_spec_from_conversation(msgs, attached_text=attached_text)
        # Stamp every AI-drafted row with attribution.
        # The 'original' snapshot for requirements is captured in the workspace
        # render once owner_name is resolved (see below). Scoring criteria are
        # snapshotted here since their fields are stable at generation time.
        now_iso = datetime.now(tz=timezone.utc).isoformat()
        for r in spec.get("requirements", []):
            r["last_edited_by"] = "AI"
            r["last_edited_at"] = now_iso
        for c in spec.get("scoring_criteria", []):
            c["last_edited_by"] = "AI"
            c["last_edited_at"] = now_iso
            c["original"] = {
                "pillar":    c.get("pillar"),
                "criterion": c.get("criterion"),
                "weight":    c.get("weight"),
                "scorer":    c.get("scorer"),
            }
        st.session_state.sb_spec = spec
        st.session_state.sb_stage = 4
        st.session_state.sb_generating = False
        if spec.get("scoring_criteria"):
            st.session_state["sm_criteria"] = [dict(c) for c in spec["scoring_criteria"]]
            for k in list(st.session_state.keys()):
                if k.startswith("sm_sent_") or k.startswith("sm_drafted_"):
                    st.session_state[k] = False
        _push_activity(
            f"generated spec for '{spec['category']}' "
            f"({len(spec['requirements'])} requirements, "
            f"{len(spec.get('scoring_criteria', []))} scoring criteria)",
            _active_user(),
            icon="🤖",
        )
        st.rerun()
        return

    if st.session_state.sb_stage < 3:
        if prompt := st.chat_input("Type your answer…"):
            stage = st.session_state.sb_stage
            msgs.append({"role": "user", "content": prompt})

            # When the user answers stage 0 (the category), generate the
            # category-specific clarifying questions before composing the
            # bot's reply.
            if stage == 0:
                cat = prompt.strip()
                st.session_state.sb_category_text = cat
                with st.spinner(f"Tailoring questions to '{cat}'…"):
                    st.session_state.sb_dynamic_questions = (
                        agents.generate_clarifying_questions(cat)
                    )

            reply, _ = _bot_reply_for_stage(stage, prompt)
            msgs.append({"role": "assistant", "content": reply})
            st.session_state.sb_stage = stage + 1
            st.rerun()


# ---------------------------------------------------------------------------
# Workspace phase: spec table + collaboration
# ---------------------------------------------------------------------------

def _owner_name_for_row(row_owner_role: str) -> str:
    """Map LLM owner role to a team member name. Fallbacks if no match."""
    u = team.first_user_for_role(row_owner_role)
    if u:
        return u["name"]
    members = team.get_team()
    return members[0]["name"] if members else "Unassigned"


def _render_workspace():
    spec = st.session_state.sb_spec
    me = _active_user()
    reqs: list[dict] = spec["requirements"]
    members = team.get_team()
    member_names = [m["name"] for m in members]

    # Migrate legacy rows: owner stored as role -> resolve once to a name.
    # Also backfill attribution for any spec generated before this feature.
    for r in reqs:
        owner_val = r.get("owner", "")
        if owner_val in ("procurement", "business", "sme"):
            r["owner_name"] = _owner_name_for_row(owner_val)
        elif owner_val and owner_val not in member_names:
            r["owner_name"] = _owner_name_for_row(owner_val)
        r.setdefault("owner_name", _owner_name_for_row(r.get("owner", "business")))
        r.setdefault("last_edited_by", "AI")
        r.setdefault("last_edited_at", datetime.now(tz=timezone.utc).isoformat())
        # Snapshot the AI draft if it wasn't captured at generation time
        if "original" not in r:
            r["original"] = {
                "section": r.get("section"),
                "title": r.get("title"),
                "description": r.get("description"),
                "priority": r.get("priority"),
                "owner_name": r.get("owner_name"),
                "status": r.get("status", "Draft"),
            }

    # Summary header
    st.markdown(f"**{spec['category']}** — {spec['summary']}")

    # ---- Product Context card ----
    spec.setdefault("context", {
        "brand_preference": "Open to alternatives",
        "geography": "UK",
        "quantity": "TBC",
        "budget": "TBC",
        "timeline": "TBC",
    })
    ctx = spec["context"]
    with st.expander("📦  Product context — brand, geography, quantity, budget, timeline", expanded=True):
        r1 = st.columns(3)
        ctx["brand_preference"] = r1[0].text_input(
            "Brand preference", value=ctx.get("brand_preference", ""), key="sb_ctx_brand",
        )
        ctx["geography"] = r1[1].text_input(
            "Geography / deployment", value=ctx.get("geography", ""), key="sb_ctx_geo",
        )
        ctx["quantity"] = r1[2].text_input(
            "Quantity / volume", value=ctx.get("quantity", ""), key="sb_ctx_qty",
        )
        r2 = st.columns(3)
        ctx["budget"] = r2[0].text_input(
            "Target budget", value=ctx.get("budget", ""), key="sb_ctx_budget",
        )
        ctx["timeline"] = r2[1].text_input(
            "Timeline", value=ctx.get("timeline", ""), key="sb_ctx_timeline",
        )
        r2[2].caption("*Edit any field — these inform the spec scope.*")

    # KPI strip
    n_total = len(reqs)
    n_mine = sum(1 for r in reqs if r.get("owner_name") == me["name"])
    n_approved = sum(1 for r in reqs if r.get("status", "").lower() == "approved")
    n_draft = n_total - n_approved

    m1, m2, m3, m4 = st.columns(4)
    m1.metric("Requirements", n_total)
    m2.metric("Assigned to you", n_mine)
    m3.metric("Approved", n_approved)
    m4.metric("Draft / Review", n_draft)

    st.write("")

    col_table, col_side = st.columns([2.6, 1])

    # ---- LEFT: spec table ----
    with col_table:
        st.markdown("##### Specification table")
        st.caption("Click any cell to edit. Rows are taller so descriptions stay readable.")

        def _edit_badge(r: dict) -> str:
            by = r.get("last_edited_by", "AI")
            at = r.get("last_edited_at", "")
            if by == "AI":
                return "🤖 AI · drafted"
            return f"✏️ {by.split()[0]} · {_human_time(at)}" if at else f"✏️ {by.split()[0]}"

        df = pd.DataFrame([
            {
                "ID": r["id"],
                "Section": r.get("section", "Technical"),
                "Requirement": r.get("title", ""),
                "Target Specification / Value": r.get("description", ""),
                "Priority": r.get("priority", "Must"),
                "Owner": r["owner_name"],
                "Status": r.get("status", "Draft").capitalize(),
                "Comments": r.get("comments", ""),
                "Last edit": _edit_badge(r),
            }
            for r in reqs
        ])

        owner_options = member_names if member_names else ["Unassigned"]

        col_cfg = {
            "ID": st.column_config.TextColumn("ID", width="small", disabled=True),
            "Section": st.column_config.SelectboxColumn(
                "Section", width="small",
                options=["Technical", "Commercial", "Legal", "Operational", "ESG"],
            ),
            "Requirement": st.column_config.TextColumn("Requirement", width="medium"),
            "Target Specification / Value": st.column_config.TextColumn(
                "Target Specification / Value", width="large",
            ),
            "Priority": st.column_config.SelectboxColumn(
                "Priority", width="small",
                options=["Must", "Should", "Could"],
            ),
            "Owner": st.column_config.SelectboxColumn(
                "Owner", width="medium",
                options=owner_options,
            ),
            "Status": st.column_config.SelectboxColumn(
                "Status", width="small",
                options=["Draft", "Under Review", "Approved"],
            ),
            "Comments": st.column_config.TextColumn("Comments", width="medium"),
            "Last edit": st.column_config.TextColumn(
                "Last edit", width="medium", disabled=True,
                help="🤖 AI = first-drafted by the agent. ✏️ Name = edited by that user.",
            ),
        }

        edited = _data_editor_tall(
            df,
            col_cfg,
            num_rows="fixed",
            key="sb_spec_editor",
            height=min(700, 100 + 80 * len(df)),
        )

        # Sync edits back to session state, stamping the editor on changed rows
        now_iso = datetime.now(tz=timezone.utc).isoformat()
        changed = False
        for i, row in edited.iterrows():
            if i >= len(reqs):
                break
            r = reqs[i]
            new_status = str(row["Status"]).capitalize()
            if (
                r.get("section") != row["Section"]
                or r.get("title") != row["Requirement"]
                or r.get("description") != row["Target Specification / Value"]
                or r.get("priority") != row["Priority"]
                or r.get("owner_name") != row["Owner"]
                or r.get("status", "Draft").capitalize() != new_status
                or r.get("comments", "") != row["Comments"]
            ):
                r["section"] = row["Section"]
                r["title"] = row["Requirement"]
                r["description"] = row["Target Specification / Value"]
                r["priority"] = row["Priority"]
                r["owner_name"] = row["Owner"]
                r["status"] = new_status
                r["comments"] = row["Comments"]
                r["last_edited_by"] = me["name"]
                r["last_edited_at"] = now_iso
                changed = True

        if changed:
            _push_activity("updated the specification table", me, icon="✏️")

        btn_l, btn_r = st.columns([1, 1])
        with btn_l:
            csv = edited.to_csv(index=False)
            st.download_button(
                "⬇  Download CSV",
                data=csv,
                file_name=f"{spec['category'].replace(' ', '_')}_spec.csv",
                mime="text/csv",
                use_container_width=True,
            )
        with btn_r:
            if st.button("↺  Start a new spec", use_container_width=True):
                for k in ["sb_messages", "sb_spec", "sb_stage", "sb_activity",
                          "sb_reminder_drafted", "sb_reminder_sent", "sb_generating",
                          "sb_dynamic_questions", "sb_category_text",
                          "sb_upload_widget", "sb_extract_cache"]:
                    st.session_state.pop(k, None)
                st.rerun()

        # ---- Track Changes panel (Word-style diff vs AI draft) ----
        _SPEC_TRACK_FIELDS = (
            ("title",       "Requirement"),
            ("description", "Target specification"),
            ("section",     "Section"),
            ("priority",    "Priority"),
            ("owner_name",  "Owner"),
            ("status",      "Status"),
        )
        rows_with_changes = []
        for r in reqs:
            orig = r.get("original") or {}
            diffs = []
            for key, label in _SPEC_TRACK_FIELDS:
                old = orig.get(key)
                new = r.get(key)
                if str(old or "") != str(new or ""):
                    diffs.append((label, old, new))
            if diffs:
                rows_with_changes.append((r, diffs))

        st.write("")
        if rows_with_changes:
            with st.expander(
                f"📝  Track changes — {len(rows_with_changes)} row(s) edited since the AI draft",
                expanded=True,
            ):
                st.caption(
                    "Each entry shows what was originally drafted by the AI vs the "
                    "current version, like Word's track-changes view."
                )
                for r, diffs in rows_with_changes:
                    by = r.get("last_edited_by", "Unknown")
                    at = _human_time(r.get("last_edited_at", "")) if r.get("last_edited_at") else ""
                    actor_label = (
                        f"<b>{by.split()[0] if by != 'AI' else by}</b>"
                        f" · <span style='color:var(--muted);font-size:11px'>{at}</span>"
                    )
                    st.markdown(
                        f"<div style='border-left:3px solid var(--brand-deep-purple);"
                        f"padding:8px 14px;margin:8px 0;background:var(--surface-2);"
                        f"border-radius:6px'>"
                        f"<div style='font-weight:700;font-size:13px'>"
                        f"{r['id']} · {r.get('title','')}</div>"
                        f"<div style='font-size:11px;color:var(--muted);margin-bottom:6px'>"
                        f"edited by {actor_label}</div>"
                        + "".join(
                            f"<div style='font-size:13px;margin:4px 0'>"
                            f"<b>{label}:</b><br>"
                            f"<del style='color:#b3303f;background:#ffe9eb;"
                            f"padding:2px 6px;border-radius:3px;text-decoration:line-through;"
                            f"margin-right:6px;display:inline-block'>"
                            f"{(str(old) if old not in (None, '') else '—')}</del>"
                            f"<ins style='color:#0a6e2a;background:#dcf6e3;"
                            f"padding:2px 6px;border-radius:3px;text-decoration:none;"
                            f"display:inline-block'>"
                            f"{(str(new) if new not in (None, '') else '—')}</ins>"
                            f"</div>"
                            for (label, old, new) in diffs
                        )
                        + "</div>",
                        unsafe_allow_html=True,
                    )
        else:
            st.caption(
                "📝  *Track changes will appear here once anyone edits a row — every "
                "change is logged against its author so the team can iterate on the AI draft.*"
            )

    # ---- RIGHT: collaboration ----
    with col_side:
        st.markdown("##### Collaboration")

        # Pending rows grouped by team member (excluding me)
        my_draft = [
            r for r in reqs
            if r.get("owner_name") == me["name"]
            and r.get("status", "Draft").lower() != "approved"
        ]
        if my_draft:
            st.info(f"**{len(my_draft)} row(s) assigned to you** still need your input.")

        pending_by_user: dict[str, list[dict]] = {}
        for r in reqs:
            owner = r.get("owner_name", "")
            if not owner or owner == me["name"]:
                continue
            if r.get("status", "Draft").lower() == "approved":
                continue
            pending_by_user.setdefault(owner, []).append(r)

        if not pending_by_user:
            if len(members) <= 1:
                st.caption(
                    "Add a teammate from the sidebar (Manage team) to collaborate "
                    "and send reminders."
                )
            else:
                st.success("All teammates' rows are approved! ✓")
        else:
            for owner_name, rows in pending_by_user.items():
                target = team.user_by_name(owner_name)
                if not target:
                    continue
                target_id = target["id"]
                sent = st.session_state.sb_reminder_sent.get(target_id, False)
                drafted = st.session_state.sb_reminder_drafted.get(target_id, False)

                if sent:
                    st.success(f"✅ Reminder sent to {owner_name.split()[0]}")
                else:
                    with st.expander(
                        f"🟠 {owner_name.split()[0]} — {len(rows)} pending",
                        expanded=True,
                    ):
                        for r in rows[:5]:
                            st.markdown(f"- {r['title']}")
                        if len(rows) > 5:
                            st.caption(f"+ {len(rows) - 5} more")

                        if st.button(
                            f"📧  Draft reminder to {owner_name.split()[0]}",
                            key=f"sb_draft_{target_id}",
                            use_container_width=True,
                        ):
                            st.session_state.sb_reminder_drafted[target_id] = not drafted

                        if st.session_state.sb_reminder_drafted.get(target_id):
                            st.caption("**Pick the rows to include in the reminder:**")
                            selected_rows = []
                            for r in rows:
                                sel_key = f"sb_pick_{target_id}_{r['id']}"
                                if sel_key not in st.session_state:
                                    st.session_state[sel_key] = True
                                if st.checkbox(
                                    f"`{r['id']}` · {r['title']}",
                                    key=sel_key,
                                ):
                                    selected_rows.append(r)

                            if not selected_rows:
                                st.warning("Select at least one row to include.")
                            else:
                                rows_txt = "\n".join(
                                    f"  - {r['id']} · {r['title']}" for r in selected_rows
                                )
                                to_addr = team.centrica_email(owner_name)
                                body = (
                                    f"Hi {owner_name.split()[0]},\n\n"
                                    f"Quick reminder — the following requirement(s) "
                                    f"in the '{spec['category']}' spec need your sign-off:\n\n"
                                    f"{rows_txt}\n\n"
                                    f"Please log in to the RFx Assistant, review each row, "
                                    f"and update the Status to 'Approved' (or add comments).\n\n"
                                    f"Thanks,\n{me['name']}"
                                )
                                st.caption(f"**To:** {to_addr}")
                                st.text_area(
                                    "body", value=body, height=200,
                                    key=f"sb_body_{target_id}",
                                    label_visibility="collapsed",
                                )
                                if st.button(
                                    f"➤  Send reminder ({len(selected_rows)} row{'s' if len(selected_rows) != 1 else ''})",
                                    type="primary",
                                    key=f"sb_send_{target_id}",
                                    use_container_width=True,
                                ):
                                    with st.spinner("Sending via Outlook…"):
                                        time.sleep(0.7)
                                    st.session_state.sb_reminder_sent[target_id] = True
                                    st.session_state.sb_reminder_drafted[target_id] = False
                                    _push_activity(
                                        f"sent reminder ({len(selected_rows)} row(s)) to {owner_name} <{to_addr}>",
                                        me, icon="📧",
                                    )
                                    st.toast(f"Reminder sent to {to_addr}", icon="📧")
                                    st.rerun()

        st.divider()
        st.markdown("**Activity**")
        feed = st.session_state.sb_activity
        if not feed:
            st.caption("No activity yet.")
        for ev in feed[:10]:
            st.markdown(
                f"{ev['icon']} **{ev['actor']}** {ev['action']}  \n"
                f"<span style='color:var(--muted);font-size:11px'>"
                f"{_human_time(ev['t'])}</span>",
                unsafe_allow_html=True,
            )


# ---------------------------------------------------------------------------
# Tall data_editor wrapper (tolerates older Streamlit lacking row_height)
# ---------------------------------------------------------------------------

def _data_editor_tall(df, col_cfg, *, num_rows: str, key: str, height: int):
    """Use row_height=72 if the installed Streamlit supports it."""
    try:
        return st.data_editor(
            df,
            column_config=col_cfg,
            use_container_width=True,
            num_rows=num_rows,
            hide_index=True,
            row_height=72,
            height=height,
            key=key,
        )
    except TypeError:
        return st.data_editor(
            df,
            column_config=col_cfg,
            use_container_width=True,
            num_rows=num_rows,
            hide_index=True,
            height=height,
            key=key,
        )
