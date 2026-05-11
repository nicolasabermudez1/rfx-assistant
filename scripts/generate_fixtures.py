"""Generate the 3 mock supplier bid files (PDF, Excel, Word).

These are realistic-looking artefacts for the demo. Run once at setup time.
"""
from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.shared import RGBColor, Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from openpyxl import Workbook
from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
from openpyxl.utils import get_column_letter
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib.colors import HexColor
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate,
    Paragraph,
    Spacer,
    Table,
    TableStyle,
    PageBreak,
)

ROOT = Path(__file__).resolve().parent.parent
BIDS = ROOT / "data" / "fixtures" / "bids"
BIDS.mkdir(parents=True, exist_ok=True)

# Centrica palette (NO RED)
NAVY = "#0F2067"
MINT = "#85DB9C"
LAVENDER = "#B999F6"
PALE_LAVENDER = "#DECFFF"
PURPLE = "#9B2BF7"


# ----------------------- Aurora — PDF -----------------------

def build_aurora_pdf() -> Path:
    out = BIDS / "Aurora_Cooling_Bid_Centrica_RFP_IT-INF-DC-COOL-2026.pdf"
    doc = SimpleDocTemplate(
        str(out),
        pagesize=A4,
        leftMargin=2 * cm,
        rightMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
        title="Aurora Cooling Systems — Centrica RFP Response",
        author="Aurora Cooling Systems Ltd",
    )

    styles = getSampleStyleSheet()
    h_title = ParagraphStyle("H_Title", parent=styles["Title"], textColor=HexColor(NAVY), fontName="Helvetica-Bold", fontSize=22, leading=26)
    h_sub = ParagraphStyle("H_Sub", parent=styles["Normal"], textColor=HexColor(PURPLE), fontName="Helvetica-Oblique", fontSize=12, leading=16)
    h1 = ParagraphStyle("H1", parent=styles["Heading1"], textColor=HexColor(NAVY), fontName="Helvetica-Bold", fontSize=16, leading=20, spaceBefore=18, spaceAfter=8)
    h2 = ParagraphStyle("H2", parent=styles["Heading2"], textColor=HexColor(PURPLE), fontName="Helvetica-Bold", fontSize=12, leading=16, spaceBefore=10, spaceAfter=4)
    body = ParagraphStyle("Body", parent=styles["BodyText"], fontName="Helvetica", fontSize=10, leading=14, spaceAfter=6)

    story = []
    story.append(Paragraph("Aurora Cooling Systems Ltd", h_title))
    story.append(Paragraph("Response to Centrica RFP IT-INF-DC-COOL-2026 — Data Centre Cooling Systems", h_sub))
    story.append(Spacer(1, 0.6 * cm))

    cover_data = [
        ["Submitted to", "Priya Rai, Senior Category Manager, Centrica plc"],
        ["Submitted by", "Geoff Marriott MBE, Programme Director — Aurora Critical Programmes"],
        ["Bid reference", "AUR-CEN-2026-DC-COOL-A+B"],
        ["Date", "18 June 2026"],
        ["Clusters bid", "Cluster A (Slough East / Hams Hall / Reading West) and Cluster B (Manchester / Glasgow)"],
        ["Headline CapEx", "£21,450,000 (Clusters A+B combined)"],
        ["10-yr OpEx", "£18,200,000"],
        ["Design-point PUE", "1.22 at full IT load (UK TMY3, ASHRAE A1)"],
        ["Validity", "90 days from submission"],
    ]
    t = Table(cover_data, colWidths=[5 * cm, 11 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (0, -1), HexColor(PALE_LAVENDER)),
        ("TEXTCOLOR", (0, 0), (0, -1), HexColor(NAVY)),
        ("FONT", (0, 0), (0, -1), "Helvetica-Bold", 10),
        ("FONT", (1, 0), (1, -1), "Helvetica", 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("BOX", (0, 0), (-1, -1), 0.5, HexColor(NAVY)),
        ("LEFTPADDING", (0, 0), (-1, -1), 8),
        ("RIGHTPADDING", (0, 0), (-1, -1), 8),
        ("TOPPADDING", (0, 0), (-1, -1), 6),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
    ]))
    story.append(t)
    story.append(PageBreak())

    # 1. Company
    story.append(Paragraph("1. Company & Capability", h1))
    story.append(Paragraph(
        "Aurora Cooling Systems Ltd (LSE: AURC) is a UK-listed critical-cooling specialist established in 1998 and "
        "headquartered in Reading. We employ 612 people across the UK, with engineering hubs in Reading, Birmingham, "
        "and Warrington. Our installed base in the UK exceeds 230 MW of critical cooling load.", body))
    story.append(Paragraph("1.1 Recent comparable installations", h2))
    ref_data = [
        ["Client", "Site", "IT load", "Year", "Year-1 PUE", "Year-3 PUE"],
        ["HSBC", "Hayes (London)", "6.4 MW", "2024", "1.23", "n/a"],
        ["NatWest Group", "Edinburgh South Gyle", "5.1 MW", "2023", "1.24", "1.27"],
        ["Department for Work & Pensions", "Telford", "4.2 MW", "2025", "1.21", "n/a"],
        ["BT", "Adastral Park (Ipswich)", "8.3 MW", "2022", "1.26", "1.28"],
    ]
    t = Table(ref_data, colWidths=[3.5 * cm, 4.0 * cm, 2.0 * cm, 1.5 * cm, 2.0 * cm, 2.0 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor(NAVY)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 9),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("BOX", (0, 0), (-1, -1), 0.5, HexColor(NAVY)),
        ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, HexColor(PALE_LAVENDER)]),
    ]))
    story.append(t)
    story.append(Paragraph("1.2 Certifications", h2))
    story.append(Paragraph(
        "ISO 14001:2015, ISO 45001:2018, ISO 27001:2022, and ISO 50001:2018 — all current to October 2027 and audited annually by BSI. "
        "Modern Slavery Act 2015 statement filed for FY25 with full tier-2 disclosure for chiller compressors (Daikin Japan), "
        "refrigerant supply (Honeywell US), and drive bearings (Schaeffler Germany).", body))

    # 2. Technical
    story.append(PageBreak())
    story.append(Paragraph("2. Technical Solution", h1))
    story.append(Paragraph(
        "Aurora proposes a chilled-water topology with 2N redundant 2.4 MW magnetic-bearing chillers per site, water-side "
        "economisers active above 7°C ambient (4,680 hrs/year against UK Met Office TMY3), and Aurora Q-series CRAH units "
        "with EC fans in N+1 configuration. The result is a design-point PUE of 1.22 at full IT load against the ASHRAE A1 "
        "envelope.", body))
    story.append(Paragraph("2.1 Refrigerant strategy", h2))
    story.append(Paragraph(
        "Our chillers are launched on R513A (GWP 631), well within the F-Gas Regulation threshold. Our committed phase-down: "
        "R513A → R454B (Q1 2027) → R290 propane (Q4 2029). Phase-down is fully cost-borne by Aurora — Centrica will not see a "
        "refrigerant change-out invoice during the contract term.", body))
    story.append(Paragraph("2.2 AI setpoint optimisation — Aurora.IQ", h2))
    story.append(Paragraph(
        "Our Aurora.IQ optimisation platform is currently in extended beta with HSBC Hayes (production go-live confirmed for "
        "September 2026). Centrica's first cluster (Hams Hall, October 2026) will be the second production deployment. We "
        "anticipate 3–5% additional energy savings on top of the headline PUE. We do NOT seek innovation-tied service credits "
        "until the platform has 12 months of production data.", body))

    # 3. Commercial
    story.append(PageBreak())
    story.append(Paragraph("3. Commercial", h1))
    capex_data = [
        ["Sub-system", "Cluster A (£)", "Cluster B (£)", "Total (£)"],
        ["Chiller plant",    "4,650,000", "3,820,000", "8,470,000"],
        ["CRAH units",       "1,920,000", "1,580,000", "3,500,000"],
        ["Pumps + pipework", "1,140,000",   "920,000", "2,060,000"],
        ["Controls + BMS",     "880,000",   "740,000", "1,620,000"],
        ["Free cooling kit",   "920,000",   "760,000", "1,680,000"],
        ["Commissioning",      "560,000",   "470,000", "1,030,000"],
        ["Training + handover","420,000",   "350,000",   "770,000"],
        ["Project mgmt",       "920,000",   "700,000", "1,620,000"],
        ["Contingency 5%",     "770,000",   "660,000", "1,430,000"],
        ["TOTAL",           "12,180,000", "9,270,000", "21,450,000"],
    ]
    t = Table(capex_data, colWidths=[4.5 * cm, 3.5 * cm, 3.5 * cm, 3.5 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor(NAVY)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("BACKGROUND", (0, -1), (-1, -1), HexColor(MINT)),
        ("FONT", (0, -1), (-1, -1), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -2), "Helvetica", 9),
        ("ALIGN", (1, 0), (-1, -1), "RIGHT"),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("BOX", (0, 0), (-1, -1), 0.5, HexColor(NAVY)),
    ]))
    story.append(t)
    story.append(Paragraph(
        "Payment terms requested: <b>Net 45 from valid invoice</b>. We note this is a deviation from Centrica's standard Net 60 "
        "and have logged it in the Deviation Register. Liability cap requested: <b>100% of charges in the prior 12 months</b> "
        "(Centrica standard is 125%). Both are open to negotiation.", body))

    # 4. Service & resilience
    story.append(PageBreak())
    story.append(Paragraph("4. Service Levels & Resilience", h1))
    story.append(Paragraph(
        "Aurora operates a 24/7/365 critical-response engineering desk from Reading with a deployed bench of 84 critical-cooling "
        "engineers across the UK: 38 South (Reading hub), 24 Midlands (Birmingham hub), 14 North-West (Warrington hub), 6 "
        "Yorkshire/Humber, 2 Scotland (Glasgow). Cluster A and Cluster B can be served from this footprint. "
        "Cluster C (Cardiff) was excluded from this bid because we cannot today commit to a 4-hour response from our Bristol team "
        "during peak weekend periods.", body))
    sla_data = [
        ["SLA", "Centrica target", "Aurora 24-month evidence"],
        ["4-hr critical response", "≥ 98%", "98.6%"],
        ["8-hr critical fix",      "≥ 95%", "97.1%"],
        ["Mean fix time (critical)", "≤ 4h",  "3h 28m"],
        ["UK spare-parts pool",    "—",     "£3.8m held across 3 warehouses"],
    ]
    t = Table(sla_data, colWidths=[5.5 * cm, 4.5 * cm, 5.5 * cm])
    t.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, 0), HexColor(NAVY)),
        ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
        ("FONT", (0, 0), (-1, 0), "Helvetica-Bold", 9),
        ("FONT", (0, 1), (-1, -1), "Helvetica", 9),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.lightgrey),
        ("BOX", (0, 0), (-1, -1), 0.5, HexColor(NAVY)),
    ]))
    story.append(t)

    # 5. Sustainability
    story.append(PageBreak())
    story.append(Paragraph("5. Sustainability & Carbon", h1))
    story.append(Paragraph(
        "Aurora's Scope 1 and 2 emissions for FY25 totalled 14,200 tCO2e, externally assured to ISAE 3410 (Limited) by KPMG. Our "
        "Science-Based Targets initiative (SBTi) submission was validated against the 1.5°C pathway in March 2024, with a Group "
        "net zero commitment of 2040.", body))
    story.append(Paragraph(
        "For Cluster A the lifecycle carbon assessment returns 6,820 kgCO2e per MW of IT load over 15 years (embodied + operational, "
        "UK grid factor projected to 2041). Make-up water target is 1.4 L/kWh of IT load — better than Centrica's 1.6 L/kWh covenant.", body))

    # 6. Implementation
    story.append(PageBreak())
    story.append(Paragraph("6. Implementation Plan", h1))
    story.append(Paragraph(
        "Programme Director: <b>Geoff Marriott MBE</b>, 22 years critical-cooling delivery (HSBC, NatWest, DWP references above). "
        "Methodology: PRINCE2 hardened against Aurora's internal critical-environments playbook. Hams Hall mobilisation date of "
        "1 October 2026 is confirmed without caveat. Slough East 18-month phased delivery aligned to IT fit-out is on the "
        "critical path of the wider Centrica programme.", body))

    # 7. Risk
    story.append(Paragraph("7. Risk & Compliance", h1))
    story.append(Paragraph(
        "F-Gas phase-down plan stated above. IEC 62443 SL-2 alignment confirmed; OT/IT segregation via dedicated VLAN with no "
        "internet ingress. Patching cadence ≤ 30 days for security patches once released by the controls vendor (Schneider EcoStruxure). "
        "Annual third-party penetration testing by Pen Test Partners; results disclosed to Centrica.", body))

    # 8. Social value
    story.append(Paragraph("8. Social Value", h1))
    story.append(Paragraph(
        "127 UK apprentices in the last 3 years (target: 165 over the FY26–28 contract window). 35% of Cluster A spend committed "
        "to local sub-contractors within the M4 corridor. Continuing partnership with the University of Reading engineering "
        "faculty (Aurora-funded chair in critical environments).", body))

    story.append(PageBreak())
    story.append(Paragraph("Signature page", h1))
    story.append(Spacer(1, 1.5 * cm))
    sig = [
        ["", "Aurora Cooling Systems Ltd", ""],
        ["", "", ""],
        ["Name:", "Geoff Marriott MBE", ""],
        ["Position:", "Programme Director — Critical Programmes", ""],
        ["Date:", "18 June 2026", ""],
        ["Signature:", "(electronically signed via DocuSign envelope #AUR-2026-0618-114)", ""],
    ]
    t = Table(sig, colWidths=[3 * cm, 9 * cm, 4 * cm])
    t.setStyle(TableStyle([
        ("FONT", (0, 0), (-1, -1), "Helvetica", 10),
        ("FONT", (0, 2), (0, -1), "Helvetica-Bold", 10),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
    ]))
    story.append(t)

    doc.build(story)
    print(f"  -> {out.name}")
    return out


# ----------------------- Helios — Excel -----------------------

def build_helios_xlsx() -> Path:
    out = BIDS / "Helios_Centrica_Cooling_RFP_Response.xlsx"
    wb = Workbook()

    navy = PatternFill("solid", fgColor=NAVY[1:])
    purple = PatternFill("solid", fgColor=PURPLE[1:])
    mint = PatternFill("solid", fgColor=MINT[1:])
    pale_lav = PatternFill("solid", fgColor=PALE_LAVENDER[1:])
    white_font = Font(color="FFFFFF", bold=True, name="Calibri", size=11)
    bold_navy = Font(color=NAVY[1:], bold=True, name="Calibri", size=11)
    body_font = Font(name="Calibri", size=10)
    title_font = Font(color=NAVY[1:], bold=True, name="Arial", size=20)
    thin = Side(border_style="thin", color="BBBBBB")
    border = Border(top=thin, left=thin, right=thin, bottom=thin)

    def style_header(row_cells):
        for c in row_cells:
            c.fill = navy
            c.font = white_font
            c.alignment = Alignment(horizontal="center", vertical="center", wrap_text=True)
            c.border = border

    def write_table(ws, start_row, headers, rows, col_widths=None):
        for i, h in enumerate(headers, start=1):
            ws.cell(row=start_row, column=i, value=h)
        style_header([ws.cell(row=start_row, column=i+1) for i in range(len(headers))])
        for r_idx, row in enumerate(rows, start=start_row + 1):
            for c_idx, v in enumerate(row, start=1):
                cell = ws.cell(row=r_idx, column=c_idx, value=v)
                cell.font = body_font
                cell.alignment = Alignment(vertical="top", wrap_text=True)
                cell.border = border
                if r_idx % 2 == 0:
                    cell.fill = pale_lav
        if col_widths:
            for i, w in enumerate(col_widths, start=1):
                ws.column_dimensions[get_column_letter(i)].width = w

    # ---- Summary ----
    ws = wb.active
    ws.title = "Summary"
    ws["A1"] = "Helios Critical Infrastructure plc"
    ws["A1"].font = title_font
    ws["A2"] = "Response to Centrica RFP IT-INF-DC-COOL-2026"
    ws["A2"].font = Font(italic=True, color=PURPLE[1:], size=12)
    ws.row_dimensions[1].height = 28

    summary = [
        ("Bid reference",     "HEL-CEN-2026-DC-COOL-ABC"),
        ("Submitted to",      "Priya Rai, Senior Category Manager, Centrica plc"),
        ("Submitted by",      "Sarah Mensah, P.E., Programme Director — Helios UK"),
        ("Date",              "19 June 2026"),
        ("Clusters bid",      "All three (A, B, C)"),
        ("Headline CapEx",    "£24,820,000 (Clusters A+B+C combined)"),
        ("10-yr OpEx",        "£16,400,000 (lowest in field)"),
        ("Design-point PUE",  "1.19 at full IT load"),
        ("Free cooling hrs/yr","4,920 (UK TMY3)"),
        ("Refrigerant",       "R1234ze (GWP 7)"),
        ("Tier",              "Uptime Tier IV (Centrica required Tier III+)"),
        ("UK 24/7 engineers", "196 across Slough, Manchester, Glasgow, Cardiff, Newcastle"),
        ("Boilerplate deviations", "ZERO"),
        ("Validity",          "150 days from submission"),
    ]
    for i, (k, v) in enumerate(summary, start=4):
        ws.cell(row=i, column=2, value=k).font = bold_navy
        ws.cell(row=i, column=2).fill = pale_lav
        ws.cell(row=i, column=2).border = border
        ws.cell(row=i, column=3, value=v).font = body_font
        ws.cell(row=i, column=3).border = border
    ws.column_dimensions["A"].width = 2
    ws.column_dimensions["B"].width = 28
    ws.column_dimensions["C"].width = 75

    # ---- Pricing_CapEx ----
    ws = wb.create_sheet("Pricing_CapEx")
    ws["A1"] = "CapEx breakdown — by cluster, by sub-system (£)"
    ws["A1"].font = title_font
    write_table(
        ws, 3,
        ["Sub-system", "Cluster A (£)", "Cluster B (£)", "Cluster C (£)", "Total (£)"],
        [
            ["Chiller plant",      4900000, 3950000, 2890000, 11740000],
            ["CRAH units",         2100000, 1720000, 1240000,  5060000],
            ["Pumps + pipework",   1280000, 1040000,  760000,  3080000],
            ["Controls + BMS",      990000,  820000,  580000,  2390000],
            ["Free cooling kit",    980000,  790000,  590000,  2360000],
            ["Commissioning",       620000,  510000,  370000,  1500000],
            ["Training + handover", 460000,  380000,  280000,  1120000],
            ["Project mgmt",       1020000,  790000,  580000,  2390000],
            ["Contingency 4%",      820000,  680000,  500000,  2000000],
            ["Sub-total",         13170000, 10680000, 7790000, 31640000],
            ["Volume rebate (-21.5%)", -2832000, -2296300, -1675000, -6803300],
            ["TOTAL after rebate",  10337000,  8383700,  6115000, 24820000],
        ],
        col_widths=[28, 16, 16, 16, 16],
    )
    ws["B22"] = "Note: Volume rebate triggered by all-three-cluster award. If Centrica awards two clusters only, rebate falls to 12%."
    ws["B22"].font = Font(italic=True, color=PURPLE[1:], size=10)

    # ---- Pricing_OpEx ----
    ws = wb.create_sheet("Pricing_OpEx")
    ws["A1"] = "10-yr OpEx model — preventive + reactive + consumables (£)"
    ws["A1"].font = title_font
    write_table(
        ws, 3,
        ["Year", "Preventive", "Reactive", "Consumables", "Software lic.", "Spares pool", "Total"],
        [
            ["FY27", 1100000, 380000, 290000, 84000, 120000, 1974000],
            ["FY28", 1130000, 390000, 295000, 84000, 120000, 2019000],
            ["FY29", 1160000, 400000, 300000, 88000, 125000, 2073000],
            ["FY30", 1180000, 420000, 305000, 88000, 125000, 2118000],
            ["FY31", 1200000, 430000, 310000, 92000, 125000, 2157000],
            ["FY32", 1220000, 440000, 315000, 92000, 125000, 2192000],
            ["FY33", 1240000, 450000, 320000, 96000, 130000, 2236000],
            ["FY34", 1260000, 460000, 325000, 96000, 130000, 2271000],
            ["FY35", 1280000, 470000, 330000, 100000, 130000, 2310000],
            ["FY36", 1300000, 480000, 335000, 100000, 130000, 2345000],
            ["TOTAL", 12070000, 4320000, 3125000, 920000, 1260000, 16400000],
        ],
        col_widths=[8, 14, 12, 14, 14, 13, 14],
    )

    # ---- Technical ----
    ws = wb.create_sheet("Technical")
    ws["A1"] = "Technical solution — Helios Vector-X"
    ws["A1"].font = title_font
    rows = [
        ["Cooling architecture",
         "Helios Vector-X chilled-water plant — 2× 3.0 MW magnetic-bearing chillers per site at N+2 redundancy. Adaptive 3-pipe flow arrangement. Water-side economiser engaged ≥9°C ambient. Helios Q4-X CRAH units with EC-fan plenum, N+1."],
        ["Design-point PUE",
         "1.19 at full IT load (UK TMY3, ASHRAE A1 envelope). Sub-PUE breakdown: chiller 0.10, pumps 0.04, CRAH 0.04, humidification 0.01."],
        ["Refrigerant",
         "R1234ze (GWP 7). Already below F-Gas 2030 cap of GWP 150. No phase-down event during contract term."],
        ["Free cooling hours",
         "4,920 hrs/yr based on UK Met Office TMY3, weighted average across the four new sites."],
        ["Redundancy posture",
         "Tier IV — N+2 chiller, N+1 CRAH, N+1 pumps, dual-feed BMS. Exceeds Centrica's Tier III+ floor."],
        ["Controls platform",
         "Helios SmartPath 4.2 with AI setpoint optimisation. 41 production EU deployments. Measured 4–7% energy saving over fixed setpoint baseline."],
        ["Immersion readiness",
         "Two-phase immersion piloting at Helios Dublin since 2024. Centrica option to retrofit any cluster post-2028."],
        ["BMS integration",
         "Native Helios SmartPath, BACnet/IP gateway to Centrica's existing Honeywell front-end. Migration path documented."],
    ]
    write_table(ws, 3, ["Topic", "Helios position"], rows, col_widths=[22, 100])

    # ---- SLAs ----
    ws = wb.create_sheet("SLAs")
    ws["A1"] = "Service levels — 24-month evidence pack"
    ws["A1"].font = title_font
    write_table(
        ws, 3,
        ["SLA", "Centrica target", "Helios 24-month evidence", "Notes"],
        [
            ["4-hr critical response", "≥ 98.0%", "99.4%", "Across 11 comparable critical sites"],
            ["8-hr critical fix",      "≥ 95.0%", "98.1%", "Across 11 comparable critical sites"],
            ["Mean response time",     "≤ 1h",    "0h 42m", "—"],
            ["Mean fix time",          "≤ 4h",    "2h 51m", "—"],
            ["UK 24/7 engineers",      "—",       "196",    "Slough 58 / Manchester 44 / Glasgow 32 / Cardiff 28 / Newcastle 34"],
            ["UK spares pool value",   "—",       "£11.2m", "5 warehouses, 24/7 dispatch"],
        ],
        col_widths=[24, 16, 26, 40],
    )

    # ---- Sustainability ----
    ws = wb.create_sheet("Sustainability")
    ws["A1"] = "Sustainability & carbon"
    ws["A1"].font = title_font
    write_table(
        ws, 3,
        ["Metric", "FY24", "FY25", "Notes"],
        [
            ["Scope 1 (tCO2e)",                  4200,  3950,  "ISAE 3410 limited assurance — Deloitte"],
            ["Scope 2 (tCO2e, market-based)",   28100, 24450,  "—"],
            ["Scope 3 (tCO2e, screening)",     162000, 158000, "Cat. 1, 4, 11 dominant"],
            ["SBTi status",                      "Validated", "Validated", "1.5°C pathway"],
            ["Net zero target year",             2035, 2035,   "5 years ahead of category covenant"],
            ["Lifecycle carbon (kgCO2e/MW/15yr)", 6090, 5940,  "Best in field"],
            ["Make-up water (L/kWh IT load)",    1.34, 1.30,   "Best in field — relevant to Cluster C Welsh catchment"],
            ["Drift eliminator",                "H-class", "H-class", "Above Centrica spec"],
        ],
        col_widths=[35, 12, 12, 50],
    )

    # ---- Risk ----
    ws = wb.create_sheet("Risk")
    ws["A1"] = "Risk & compliance"
    ws["A1"].font = title_font
    rows = [
        ["F-Gas (Reg EU 517/2014, retained UK)", "Compliant. R1234ze (GWP 7) already below 2030 cap of GWP 150. No phase-down required during contract."],
        ["IEC 62443 alignment", "SL-3 (above Centrica required SL-2). Annual NCC Group penetration test, results shared with Centrica TPRM."],
        ["Tier-2 supply chain", "Compressor (Daikin), refrigerant (Honeywell), drives (ABB), controls silicon (NXP). Full disclosure provided in Annex C."],
        ["Cyber patching cadence", "≤ 14 days for security patches once released by controls vendor."],
        ["Modern Slavery Act 2015", "Statement filed FY25, full tier-2 disclosure included."],
        ["M&A disclosure", "Trade-press reports of rumoured Stulz acquisition. Helios will accept a contractual M&A change-of-control clause aligned to Centrica boilerplate §10."],
    ]
    write_table(ws, 3, ["Topic", "Helios position"], rows, col_widths=[28, 100])

    # ---- Deviations ----
    ws = wb.create_sheet("Deviations")
    ws["A1"] = "Boilerplate deviations register"
    ws["A1"].font = title_font
    ws["A3"] = "ZERO deviations. Helios accepts the Centrica boilerplate verbatim including liability cap (125%), payment terms (Net 60), step-in (20 BD), F-Gas covenant (2030 GWP ≤ 150), and PI insurance £25m."
    ws["A3"].font = Font(bold=True, color=NAVY[1:], size=12)
    ws["A3"].fill = mint
    ws["A3"].alignment = Alignment(wrap_text=True, vertical="top")
    ws.row_dimensions[3].height = 60
    ws.merge_cells("A3:E3")
    ws.column_dimensions["A"].width = 30

    wb.save(out)
    print(f"  -> {out.name}")
    return out


# ----------------------- NorthernAir — Word -----------------------

def build_northair_docx() -> Path:
    out = BIDS / "NorthernAir_Centrica_Bid.docx"
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    def heading(text, level=1, color=NAVY):
        p = doc.add_heading(text, level=level)
        for run in p.runs:
            run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
            run.font.name = "Arial"
        return p

    def para(text, italic=False, bold=False, color=None):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.italic = italic
        run.bold = bold
        if color:
            run.font.color.rgb = RGBColor.from_string(color.lstrip("#"))
        return p

    def shaded_table(rows, header=True):
        tbl = doc.add_table(rows=len(rows), cols=len(rows[0]))
        tbl.style = "Light Grid Accent 1"
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                cell = tbl.rows[r].cells[c]
                cell.text = str(val)
                if header and r == 0:
                    for run in cell.paragraphs[0].runs:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor.from_string(NAVY.lstrip("#"))
        return tbl

    # ----- Cover -----
    p = doc.add_paragraph()
    run = p.add_run("NorthernAir Solutions GmbH")
    run.font.size = Pt(28); run.font.bold = True; run.font.name = "Arial"
    run.font.color.rgb = RGBColor.from_string(NAVY.lstrip("#"))
    para("Response to Centrica RFP IT-INF-DC-COOL-2026", italic=True, color=PURPLE)
    para("Submitted by NorthernAir UK Ltd, the UK trading entity of NorthernAir Solutions GmbH (Stuttgart)", italic=True, color="555555")

    shaded_table([
        ["Bid reference",    "NA-CEN-2026-DC-COOL-A+C"],
        ["Submitted to",     "Priya Rai, Senior Category Manager, Centrica plc"],
        ["Submitted by",     "Helmut Bauer, Group Programme Director (UK time 30%)"],
        ["UK contact",       "James Henderson, MD NorthernAir UK Ltd"],
        ["Date",             "19 June 2026"],
        ["Clusters bid",     "Cluster A (Slough East / Hams Hall / Reading West) and Cluster C (Cardiff Imperial Park)"],
        ["Headline CapEx",   "£19,980,000 — lowest in field"],
        ["10-yr OpEx",       "£19,850,000"],
        ["Design-point PUE", "1.24 — meets Centrica 1.25 ceiling"],
        ["Validity",         "60 days from submission"],
    ])

    doc.add_page_break()

    # ----- Section 1 -----
    heading("1. Company & Capability", 1)
    para(
        "NorthernAir Solutions GmbH is a German Mittelstand company founded 1962, headquartered in Stuttgart, and "
        "privately owned by the Vogt family. We employ 247 people in our UK subsidiary NorthernAir UK Ltd, with offices in "
        "Coventry (engineering hub) and Bristol (sales + service)."
    )
    heading("1.1 UK references", 2, color=PURPLE)
    shaded_table([
        ["Client",       "Site",                "IT load", "Year", "Year-1 PUE"],
        ["BNP Paribas",  "London Bishopsgate",  "3.2 MW", "2024", "1.23"],
        ["Bayer UK",     "Reading West",        "2.4 MW", "2023", "1.25"],
        ["Roche",        "Welwyn",              "4.0 MW", "2025", "1.22"],
    ])
    heading("1.2 Certifications", 2, color=PURPLE)
    para("ISO 14001:2015, ISO 45001:2018, ISO 27001:2022 all current. ISO 50001:2018 in progress, target Q4 2026.")

    # ----- Section 2 -----
    heading("2. Technical Solution", 1)
    para(
        "We propose our flagship NorthernAir Adiabatic-Plus indirect adiabatic system with chilled-water trim. Adiabatic-Plus "
        "is genuinely best-in-class for water stewardship — make-up water target 1.1 L/kWh of IT load, materially below Centrica's "
        "1.6 L/kWh covenant. This is particularly relevant for Cluster C (Cardiff Imperial Park) given the Welsh Water 'moderate "
        "stress' classification of the catchment."
    )
    para(
        "Design-point PUE: 1.24 at full IT load (UK TMY3, ASHRAE A1) — meets Centrica's 1.25 ceiling without margin. Free-cooling "
        "hours: 4,450 per year, above Centrica's 4,200 floor (note: figure includes 6% safety-margin derating from supplier-side; "
        "raw model output is 4,710 hrs/yr)."
    )
    para(
        "Redundancy: N+1 across chillers, CRAH, pumps, controls — meets Centrica spec, no headroom above. Refrigerant: R454B "
        "(GWP 466). AI setpoint optimisation: on the NorthernAir Group roadmap for delivery 2027 H2; not deployed at any "
        "production site today."
    )

    # ----- Section 3 -----
    doc.add_page_break()
    heading("3. Commercial", 1)
    shaded_table([
        ["Sub-system", "Cluster A (£)", "Cluster C (£)", "Total (£)"],
        ["Chiller plant",      "4,180,000", "1,840,000",  "6,020,000"],
        ["Adiabatic units",    "2,420,000", "1,070,000",  "3,490,000"],
        ["Pumps + pipework",   "1,070,000",   "470,000",  "1,540,000"],
        ["Controls + BMS",       "780,000",   "340,000",  "1,120,000"],
        ["Free cooling kit",     "880,000",   "390,000",  "1,270,000"],
        ["Commissioning",        "510,000",   "230,000",    "740,000"],
        ["Training + handover",  "390,000",   "180,000",    "570,000"],
        ["Project mgmt",         "880,000",   "430,000",  "1,310,000"],
        ["Contingency 9%",      "1,260,000",   "660,000",  "1,920,000"],
        ["TOTAL",              "12,370,000", "5,610,000", "19,980,000"],
    ])
    para(
        "Payment terms requested: Net 30 from valid invoice (deviation from Centrica's Net 60). Liability cap requested: 100% "
        "of charges in the prior 12 months (deviation from Centrica's 125%). Price validity: 60 days from submission."
    )

    # ----- Section 4 -----
    heading("4. Service Levels & Resilience", 1)
    para(
        "NorthernAir UK operates 38 critical-response engineers across Coventry (24) and Bristol (14). We confirm 24/7/365 "
        "coverage for Cluster A from this footprint."
    )
    para(
        "For Cluster C (Cardiff Imperial Park), night-shift response is sub-contracted to a Welsh Tier 1 facilities partner. "
        "The partner identity is being finalised at the time of submission and will be confirmed in the post-bid clarifications "
        "window. We acknowledge this introduces a TPRM gating step.",
        italic=True,
    )
    shaded_table([
        ["SLA",                       "Centrica target", "NorthernAir 24-mo evidence"],
        ["4-hr critical response",    "≥ 98.0%",         "96.8%"],
        ["8-hr critical fix",         "≥ 95.0%",         "92.4%"],
        ["Mean fix time (critical)",  "≤ 4h",            "4h 12m"],
        ["UK spares pool",            "—",               "£1.6m"],
    ])
    para(
        "We acknowledge our SLA evidence sits below the Centrica required floors. We have committed in our remediation plan to "
        "onboarding a further 12 UK engineers in H2 2026 if the bid is successful, which our internal modelling indicates would "
        "lift the response and fix metrics to 98.5% / 95.5% within 12 months of award.",
        italic=True,
    )

    # ----- Section 5 -----
    doc.add_page_break()
    heading("5. Sustainability & Carbon", 1)
    para(
        "NorthernAir Group reported Scope 1+2 emissions of 6,800 tCO2e for FY25 (group, not UK-only). SBTi status: committed, "
        "not yet validated — submission expected H1 2027. Net zero target year: 2040."
    )
    para(
        "Lifecycle carbon assessment for the proposed Cluster A and Cluster C designs: available on request. We acknowledge "
        "this is a §5 mandatory item in the RFP and we commit to providing it within 7 working days of bid submission.",
        italic=True,
    )
    para(
        "Make-up water target: 1.1 L/kWh of IT load — best in field. This is a genuine engineering strength of the Adiabatic-Plus "
        "platform and is directly relevant to Cluster C's Welsh Water catchment classification."
    )

    # ----- Section 6 -----
    heading("6. Implementation", 1)
    para(
        "Programme Director: Helmut Bauer (Stuttgart, UK-time 30%). On-the-ground UK lead: James Henderson (MD NorthernAir UK Ltd, "
        "based Coventry). Methodology: PRINCE2 hybrid against Group standard ProjektTakt."
    )
    para(
        "Critical-path note: chiller fabrication runs from the Stuttgart factory with a 12-week lead time. Hams Hall mobilisation "
        "of 1 October 2026 is achievable but tight; a 4-week slippage in chiller release would affect commissioning in November.",
        italic=True,
    )

    # ----- Section 7 -----
    heading("7. Risk & Compliance", 1)
    para(
        "F-Gas: R454B (GWP 466) is within today's regulatory thresholds. Group-wide phase-down to GWP < 150 refrigerants is "
        "scheduled for 2032 — we acknowledge this is two years beyond Centrica's contractual covenant of 2030 and have logged it "
        "as a deviation. We would welcome dialogue on a project-specific exemption.",
        italic=True,
    )
    para(
        "IEC 62443: alignment to SL-2 is in progress and expected to complete in Q4 2026 ahead of first commissioning. "
        "OT/IT segregation: in place at the controls platform level."
    )

    # ----- Section 8 -----
    heading("8. Social Value", 1)
    para(
        "44 UK apprentices in the last 3 years. Local sub-contractor commitment: 18% Cluster A, 24% Cluster C. STEM outreach: "
        "Coventry University engineering faculty partnership."
    )

    # ----- Deviations -----
    doc.add_page_break()
    heading("9. Deviation Register", 1)
    shaded_table([
        ["#", "RFP §", "Centrica position", "NorthernAir position", "Risk owner if accepted"],
        ["1", "§1 Liability cap",       "125%",               "100%",                "Procurement"],
        ["2", "§3 Payment terms",       "Net 60",             "Net 30",              "Finance"],
        ["3", "§4 Step-in notice",      "20 Business Days",   "14 Business Days",    "Legal"],
        ["4", "§6(d) F-Gas phase-down", "≤ 2030 GWP 150",     "≤ 2032 GWP 150",      "ESG (Rob Fenwick)"],
        ["5", "§8 PI insurance",        "≥ £25m",             "£15m",                "Risk (Aisha Bello)"],
    ])

    heading("Signature", 1)
    para("Helmut Bauer, Group Programme Director — NorthernAir Solutions GmbH", bold=True)
    para("Electronically signed via DocuSign envelope #NA-2026-0619-228, 19 June 2026.", italic=True)

    doc.save(out)
    print(f"  -> {out.name}")
    return out


# ----------------------- main -----------------------

def main():
    print("Generating supplier bid fixtures...")
    build_aurora_pdf()
    build_helios_xlsx()
    build_northair_docx()
    print("\nAll bid fixtures generated successfully.")


if __name__ == "__main__":
    main()
